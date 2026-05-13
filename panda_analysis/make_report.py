# -*- coding: utf-8 -*-
"""
make_report.py — 生成《旅游文本数据分析》课程实验报告(.docx)
================================================================
按 D:\\Desktop\\作业模版(1).docx 的章节框架,把已经跑完的全部分析结果
组装成一份学位级别 Word 报告。

用法:
    cd panda_analysis
    .venv\\Scripts\\Activate.ps1
    python make_report.py

输出: ./学号_姓名_分析报告.docx
"""

import os
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")


# ============================================================
# 一、学生信息(改这里就改了)
# ============================================================
STUDENT_ID = "202422480546"
STUDENT_NAME = "罗靖坤"
STUDENT_MAJOR = "旅游管理"
STUDENT_COLLEGE = "旅游学院"
TEACHER_NAME = "徐 欣"
SUBMIT_DATE = "2026年 5月"


OUT_DIR = Path("output")
REPORT_PATH = Path(f"{STUDENT_ID}_{STUDENT_NAME}_分析报告.docx")


# ============================================================
# 动态统计数据(从 output/ 里的 CSV 现算,避免硬编码)
# ============================================================
def _load_stats() -> dict:
    """从已有产出中计算报告里需要引用的关键数字。"""
    stats = {}
    stats["sample_n"] = 2000
    stats["n_dom"] = 1000
    stats["n_intl"] = 1000
    stats["sample_date_start"] = "2024-01-01"
    stats["sample_date_end"] = "2026-05-10"
    stats["word_freq_n"] = 2000
    stats["word_counts"] = {}

    sample_csv = Path("data/panda_sample_2000.csv")
    if sample_csv.exists():
        sample_df = pd.read_csv(sample_csv, encoding="utf-8-sig")
        stats["sample_n"] = len(sample_df)
        if "group" in sample_df.columns:
            stats["n_dom"] = int((sample_df["group"] == "国内游客").sum())
            stats["n_intl"] = int((sample_df["group"] == "国际游客").sum())
        if "publishDate" in sample_df.columns:
            dates = pd.to_datetime(sample_df["publishDate"], errors="coerce").dropna()
            if len(dates):
                stats["sample_date_start"] = dates.min().strftime("%Y-%m-%d")
                stats["sample_date_end"] = dates.max().strftime("%Y-%m-%d")

    stopwords_path = Path("stopwords.txt")
    stats["stopwords_count"] = (
        sum(1 for line in stopwords_path.read_text(encoding="utf-8").splitlines()
            if line.strip())
        if stopwords_path.exists() else 0
    )
    custom_dict_path = Path("custom_dict.txt")
    stats["custom_dict_count"] = (
        sum(1 for line in custom_dict_path.read_text(encoding="utf-8").splitlines()
            if line.strip())
        if custom_dict_path.exists() else 0
    )

    # 词频统计
    freq_csv = OUT_DIR / "1_word_freq" / "all_words_freq.csv"
    if freq_csv.exists():
        freq_df = pd.read_csv(freq_csv, encoding="utf-8-sig")
        stats["total_tokens"] = int(freq_df["count"].sum())
        stats["unique_words"] = len(freq_df)
        stats["word_counts"] = dict(zip(freq_df["word"], freq_df["count"]))
    else:
        stats["total_tokens"] = 0
        stats["unique_words"] = 0

    # 情感分析 - 机翻比例
    sent_csv = OUT_DIR / "3_sentiment" / "sentiment_results.csv"
    if sent_csv.exists():
        sent_df = pd.read_csv(sent_csv, encoding="utf-8-sig")
        stats["n_total"] = len(sent_df)
        if "is_translated" in sent_df.columns:
            n_tr = int(sent_df["is_translated"].fillna(False).astype(bool).sum())
            stats["n_translated"] = n_tr
            stats["pct_translated"] = round(n_tr / max(len(sent_df), 1) * 100, 1)
            stats["word_freq_n"] = max(len(sent_df) - n_tr, 0)
            # 国际组机翻比例
            intl = sent_df[sent_df["group"] == "国际游客"]
            n_tr_intl = int(intl["is_translated"].fillna(False).astype(bool).sum())
            stats["pct_translated_intl"] = round(
                n_tr_intl / max(len(intl), 1) * 100, 1)
        else:
            stats["n_translated"] = 0
            stats["pct_translated"] = 0.0
            stats["pct_translated_intl"] = 0.0
    else:
        stats["n_total"] = 2000
        stats["n_translated"] = 0
        stats["pct_translated"] = 0.0
        stats["pct_translated_intl"] = 0.0

    dim_csv = OUT_DIR / "4_themes" / "dimensions_table.csv"
    stats["max_dim_gap_name"] = "国内外差异最大的维度"
    stats["max_dim_gap"] = 0.0
    if dim_csv.exists():
        dim_df = pd.read_csv(dim_csv, encoding="utf-8-sig")
        needed = {"维度名称", "国内情感得分", "国际情感得分"}
        if needed.issubset(dim_df.columns) and len(dim_df):
            gaps = (dim_df["国内情感得分"].astype(float)
                    - dim_df["国际情感得分"].astype(float)).abs()
            idx = gaps.idxmax()
            stats["max_dim_gap_name"] = str(dim_df.loc[idx, "维度名称"])
            stats["max_dim_gap"] = round(float(gaps.loc[idx]), 3)

    return stats


# 模块级加载(main 里 chdir 之后再调)
STATS: dict = {}


# 图片路径快捷方式
IMG_WC_ALL          = OUT_DIR / "1_word_freq" / "all_wordcloud.png"
IMG_WC_DOM          = OUT_DIR / "1_word_freq" / "domestic_wordcloud.png"
IMG_WC_INTL         = OUT_DIR / "1_word_freq" / "intl_wordcloud.png"
IMG_COMPARISON      = OUT_DIR / "1_word_freq" / "comparison_top30.png"
IMG_DISTINCT        = OUT_DIR / "1_word_freq" / "distinctive_words.png"
IMG_POS_PIE         = OUT_DIR / "1_word_freq" / "pos_distribution.png"
IMG_NET_ALL         = OUT_DIR / "2_network"   / "network_static_all.png"
IMG_NET_COMPARE     = OUT_DIR / "2_network"   / "network_comparison.png"
IMG_SENT_DIST       = OUT_DIR / "3_sentiment" / "sentiment_distribution.png"
IMG_SENT_WEIGHTED   = OUT_DIR / "3_sentiment" / "sentiment_weighted_distribution.png"
IMG_SENT_GROUP      = OUT_DIR / "3_sentiment" / "sentiment_by_group.png"
IMG_SENT_ASPECT     = OUT_DIR / "3_sentiment" / "sentiment_by_aspect_grouped.png"
IMG_SENT_TREND      = OUT_DIR / "3_sentiment" / "sentiment_trend.png"
IMG_RADAR           = OUT_DIR / "4_themes"    / "dimensions_radar.png"


# ============================================================
# 二、底层排版辅助
# ============================================================
def set_cn_font(run, font_name="宋体", size_pt=12, bold=False, color=None):
    """为 run 设置中文字体、字号、粗体、颜色。"""
    run.font.name = font_name
    # 让中文也走同一字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text, font="宋体", size=12, bold=False,
          align=None, indent_first=True, line_spacing=1.5, color=None,
          space_before=0, space_after=0):
    """添加一段正文。indent_first=True 时首行缩进 2 字符。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if line_spacing:
        pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Pt(size * 2)  # 约 2 字符
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_cn_font(run, font_name=font, size_pt=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    """中文标题:
       level=0 → 一级章(三号黑体)
       level=1 → 二级节(小三黑体)
       level=2 → 三级目(四号黑体)
    """
    sizes = {0: 16, 1: 15, 2: 14}
    spaces = {0: (18, 12), 1: (12, 8), 2: (8, 6)}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(spaces[level][0])
    pf.space_after = Pt(spaces[level][1])
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_cn_font(run, font_name="黑体", size_pt=sizes[level], bold=True)
    # 大纲级别(让 TOC 可识别)
    pPr = p._p.get_or_add_pPr()
    outline_lvl = OxmlElement("w:outlineLvl")
    outline_lvl.set(qn("w:val"), str(level))
    pPr.append(outline_lvl)
    return p


def add_image(doc, image_path, width_cm=14, caption=None):
    """居中插入一张图片,下方添加图注。"""
    if not Path(image_path).exists():
        add_p(doc, f"[图片缺失:{image_path}]", color=(200, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.first_line_indent = Pt(0)
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run(caption)
        set_cn_font(cap_run, font_name="楷体", size_pt=10.5)


def add_table_from_df(doc, df, header_bg="2E75B6", caption=None,
                      col_widths_cm=None, font_size=10.5):
    """根据 DataFrame 生成表格(首行表头白字蓝底)。"""
    n_rows, n_cols = df.shape[0] + 1, df.shape[1]
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # 表头
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run(str(col))
        set_cn_font(run, font_name="黑体", size_pt=font_size, bold=True, color=(255, 255, 255))
        # 涂底色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_bg)
        tcPr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, val in enumerate(row.tolist()):
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            run = para.add_run(str(val))
            set_cn_font(run, font_name="宋体", size_pt=font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 列宽
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[j].width = Cm(w)

    # 图注
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.first_line_indent = Pt(0)
        cap_p.paragraph_format.space_before = Pt(2)
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run(caption)
        set_cn_font(cap_run, font_name="楷体", size_pt=10.5)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_number_footer(doc):
    """页脚:使用静态文本,避免 headless 渲染器不更新 PAGE 字段导致空页码。"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f"{STUDENT_ID}  {STUDENT_NAME}  ·  旅游文本数据分析课程实验报告")
    set_cn_font(run, "宋体", 10)


# ============================================================
# 三、内容生成
# ============================================================
def build_cover(doc):
    """封面页。"""
    # 顶部空行
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)

    # 主标题:旅游大数据理论与分析
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("旅游大数据理论与分析")
    set_cn_font(run, "黑体", 28, bold=True)

    # 课程实验
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("课程实验")
    set_cn_font(run, "黑体", 22, bold=True)

    # 实验内容
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(60)
    run = p.add_run("实验内容:旅游文本数据分析")
    set_cn_font(run, "宋体", 18, bold=True)

    # 副标题 — 论文题目
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("成都大熊猫繁育研究基地游客评论文本挖掘")
    set_cn_font(run, "黑体", 18, bold=True, color=(46, 117, 182))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(80)
    run = p.add_run("——基于国内/国际游客的跨文化对比研究")
    set_cn_font(run, "楷体", 15, color=(46, 117, 182))

    # 学生信息表
    info = [
        ("学    院:", STUDENT_COLLEGE),
        ("专    业:", STUDENT_MAJOR),
        ("学生姓名:", STUDENT_NAME),
        ("学生学号:", STUDENT_ID),
        ("指导教师:", TEACHER_NAME),
    ]
    for label, val in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"{label}    {val}")
        set_cn_font(run, "宋体", 16)

    # 日期
    for _ in range(2):
        doc.add_paragraph().paragraph_format.first_line_indent = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f"提交日期:{SUBMIT_DATE}")
    set_cn_font(run, "宋体", 16)


def build_ch1_collection(doc):
    """第一章 数据采集"""
    add_heading(doc, "一、数据采集", level=0)

    add_heading(doc, "1.1 数据来源与采集思路", level=1)
    add_p(doc,
        "本研究以成都大熊猫繁育研究基地为案例,围绕国内与国际游客的"
        "评论差异展开文本挖掘。文本数据来自国内最大的在线旅游平台——"
        "携程旅游(Ctrip)及其国际版 Trip.com 站点。两者底层共用同一"
        "套评论数据库,但通过 fromTypeText 字段区分评论来源,可自然分"
        "出"
        "「国内游客」与「国际游客」两组样本,为后续跨文化对比提供数据基础。")
    add_p(doc,
        "选择该案例的原因在于:(1) 大熊猫基地是成都最具国际知名度的"
        "城市旅游 IP,既有大量国内本地及周边游客,也有来自全球的"
        "国际游客,样本天然具有跨文化"
        "可比性;(2) 评论文本量充足、内容信息密度高,适合做词频、"
        "语义网络、情感等多维度分析;(3) 携程接口对 POI 的评论收录"
        "完整,可同时拿到 1~5 星全分布的真实评价。")

    add_heading(doc, "1.2 数据采集流程", level=1)
    add_p(doc,
        "数据采集采用 Python + requests 编写的网络爬虫,目标 API 为"
        " m.ctrip.com/restapi/soa2/13444/json/getCommentCollapseList,"
        "对应 POI(大熊猫基地)的 poiId 为 76342。整体采集流程如下:")
    steps = [
        "① 抓包定位接口:通过浏览器开发者工具(F12 → Network)在景点"
        "详情页触发评论加载,定位到 getCommentCollapseList 接口,确定"
        "其请求参数(poiId、pageIndex、starType、sortType 等)、请求"
        "头(User-Agent、cookieOrigin、UBT_VID 等)及 JSON 响应结构。",
        "② Cookie 维护:从开发者工具中复制访客 Cookie 字段(UBT_VID、"
        "GUID、_bfa 等),剔除登录态字段后写入爬虫脚本中作为请求凭证。",
        "③ 分层采样策略:鉴于携程评分严重偏正向,若按时间顺序爬取将出现"
        "「全是好评」的样本偏差。本研究按星级(1~5 星)分别爬取,"
        "确保差评/中评/好评在分析样本中都有足够数量。由于该样本不是"
        "自然总体比例,涉及整体口碑时另按清洗后总体的 group×score 分布"
        "进行加权估计。",
        "④ 断点续爬与容错:脚本启动时自动读取已抓取的 CSV、跳过已存在的 "
        "commentId,可在网络中断后从指定页码续抓;网络错误自动重试 3 次,"
        "连续 3 页空结果或失败则停止;遇到 HTTP 403/429 立即停止以规避"
        "反爬策略。",
        "⑤ 反爬规避:每页请求之间随机延时 2~4 秒,模拟正常浏览节奏;"
        "请求头沿用浏览器抓包得到的稳定参数;不并发请求,以低速率换稳定性。",
        "⑥ 数据落盘:每抓到一页就追加写入 CSV(UTF-8 with BOM),"
        "字段包括 commentId、page、content、score、publishTime、"
        "userNick、ipLocatedName、touristTypeDisplay、"
        "usefulCount、imageCount、videoCount、fromTypeText 等共 14 列。",
    ]
    for s in steps:
        add_p(doc, s)
    add_p(doc,
        "本研究原计划同步爬取马蜂窝平台 500 条评论作为补充,但马蜂窝全站"
        "受 Akamai Bot Manager 保护:所有 HTML/接口请求均返回 HTTP 202 +"
        " probe.js 指纹挑战,requests 直连无法穿过。在权衡时间成本后,"
        "本研究放弃马蜂窝部分,转为通过携程 fromTypeText 字段区分国内/"
        "国际样本,同样达到了跨文化对比的目的。")

    add_heading(doc, "1.3 数据集情况与核心定义", level=1)
    add_p(doc,
        f"本研究最终采集并清洗整理出 {STATS['sample_n']} 条有效评论,其中"
        f"国内游客 {STATS['n_dom']} 条、国际游客 {STATS['n_intl']} 条。"
        f"评论时间范围为 {STATS['sample_date_start']} 至 "
        f"{STATS['sample_date_end']},涵盖家庭亲子、情侣夫妻、朋友出游、单独旅行等多种"
        "出游类型。表 1-1 列出了最终数据集的核心字段定义。")

    add_heading(doc, "1.3.1 「国际游客」操作性定义", level=2)
    add_p(doc,
        "携程/Trip.com 接口不提供用户国籍字段,本研究以平台来源与 IP "
        "属地作为代理变量,将「国际游客」操作性定义为:通过 Trip.com "
        "国际版平台发布评论(fromTypeText = \"来自Trip.com\"),或 IP 属地"
        "为境外国家(新加坡、日本、马来西亚、泰国、美国、澳大利亚、"
        "英国、德国、荷兰、韩国等 15 国)的用户。港澳台地区因评论语言"
        "为中文原生且通过国内订单渠道(\"来自订单\")发布,归入国内组。"
        "该定义下,国际组中约 2.3% 的评论(23 条)为 IP 境外但通过"
        "\"来自订单\"渠道发布,内容检查显示其为中文母语者(可能为华侨、"
        "留学生或出差人员),其旅游消费行为仍具有跨文化参考价值,故保留"
        "在国际组中。")

    add_heading(doc, "1.3.2 抽样策略与统计口径说明", level=2)
    add_p(doc,
        "本研究的 2000 条分析样本采用评分分层抽样(stratified sampling)"
        "策略:在国内组和国际组内部,分别按 1~5 星设定配额抽取,目的是"
        "保证各评分层(尤其是低分层)都有足够样本用于组间对比和文本分析。"
        "因此,该样本的评分分布不等于携程平台上该景点的自然评分分布。")
    add_p(doc,
        "报告中涉及两套统计口径:(1)「分层样本(未加权)」——直接基于"
        " 2000 条样本计算,适用于国内/国际对比、各评分层内部特征分析;"
        "(2)「按总体评分分布加权」——以清洗后全量数据(7023 条)中"
        "每个 group×score 单元格的实际数量为权重,对分层样本的情感比例"
        "进行校正,用于估计更接近自然总体口碑的情感倾向。两套口径在"
        "报告图表中均有明确标注,读者应根据分析目的选择对应口径解读。")

    # 字段表
    fields = pd.DataFrame({
        "字段名": ["content_clean", "score", "publishDate", "group",
                  "touristTypeDisplay", "ipLocatedName", "fromTypeText", "userNick"],
        "含义": ["清洗后的评论文本(分析主字段)", "评分(1-5 星)",
                "发布日期 YYYY-MM-DD HH:MM:SS",
                "分组:国内游客 / 国际游客",
                "出游类型:家庭亲子/情侣/朋友/单独",
                "IP 属地(省份/国家)",
                "评论来源(来自订单/来自 Trip.com)",
                "用户昵称(平台已脱敏)"],
        "用途": ["词频/语义/情感分析", "情感交叉验证", "时间趋势分析",
                "跨文化对比核心字段", "人群细分", "地域分析",
                "国内/国际划分依据", "备用"],
    })
    add_table_from_df(doc, fields,
                      caption="表 1-1  最终数据集主要字段说明",
                      col_widths_cm=[4, 6, 5], font_size=10.5)


def build_ch2_preprocess(doc):
    add_heading(doc, "二、数据预处理", level=0)
    add_p(doc,
        "原始爬取数据存在时间戳格式异常、表情符号干扰、重复评论、机器"
        "水军、超短无效评论等问题,直接用于分析会显著拉低词频和语义"
        "网络的信号质量。本研究的数据预处理分两层:一是面向「干净文本」"
        "的字段级清洗;二是面向「分析输入」的分词与停用词配置。")

    add_heading(doc, "2.1 字段级清洗", level=1)
    cleaning = [
        "① 时间戳解析:将携程返回的 /Date(1778431717000+0800)/ 格式解析"
        "为标准 YYYY-MM-DD HH:MM:SS 字符串,便于后续按月度做趋势分析。",
        "② 表情与特殊字符处理:使用正则匹配剔除评论中的 emoji(包括"
        " U+1F300-U+1FAFF 和 U+2600-U+27BF 区段)、URL 链接、"
        "@用户名以及连续空白字符。",
        "③ 重复评论去重:按清洗后的文本完全相同去重,共剔除若干条复制"
        "粘贴评论。",
        "④ 极短评论删除:字数小于 10 字的评论(如「好」「不错」)"
        "信息量过低,直接删除以避免污染词频。",
        "⑤ 疑似水军过滤:同一用户(userNick)发布超过 5 条评论的判定"
        "为水军账号,该账号下所有评论一并剔除。",
        "⑥ 字段对齐:将国内版与国际版评论字段标准化对齐,group 字段"
        "基于 fromTypeText 与 IP 属地共同生成:「来自 Trip.com」或"
        "境外国别 IP → 国际游客,其余 → 国内游客。",
        "⑦ 编码与采样:输出 UTF-8 with BOM 编码的 CSV(便于 Excel 直接"
        "打开),并按 group 字段下采样,保证国内 1000 条 + 国际 1000 条"
        "的均衡样本。",
    ]
    for s in cleaning:
        add_p(doc, s)

    add_heading(doc, "2.2 分词与停用词配置", level=1)
    add_p(doc,
        "中文文本分词采用业界主流的 jieba 分词库。由于熊猫基地评论中"
        "包含大量专有名词,默认词典容易将其切碎(例如把「大熊猫繁育"
        "研究基地」拆成「大熊猫/繁育/研究/基地」、把「花花」拆为单字"
        "「花/花」),严重影响后续分析效果。为此本研究构建了"
        "三类辅助资源:")
    aux = [
        f"① 自定义词典(custom_dict.txt,{STATS['custom_dict_count']} 个词):包含「大熊猫繁育"
        "研究基地」「熊猫基地」「熊猫宝宝」「月亮产房」「太阳产房」"
        "「青年别墅」「亚成年别墅」「大熊猫博物馆」「小熊猫」"
        "「花花」「和花」「和叶」「福宝」「萌兰」「观光车」「讲解员」"
        "「熊猫基地北门」「熊猫基地南门」「亲子游」「打卡」「避雷」"
        "等熊猫领域专有术语和场景词,赋予高词频(100)以保证整体切分。",
        f"② 中文停用词表(stopwords.txt,{STATS['stopwords_count']} 个词):基于哈工大停用词表"
        "扩展,包含「的、了、是、在、和」等通用功能词,以及针对旅游"
        "评论场景补充的噪音词(如「景区、感觉、觉得、确实、就是」等)。"
        "试跑阶段发现「携程、根本、不会」被 jieba 错标为名词放过过滤,"
        "增补到停用词表后噪音得到有效控制。",
        "③ 词性过滤策略:采用 jieba.posseg 进行带词性标注的分词,仅保留"
        "名词族(n/nr/ns/nt/nz)、动词(v)和形容词(a),并对每个词"
        f"施加长度(2 ≤ len ≤ 8)、纯数字、特殊符号三重过滤。最终 2000"
        f"条评论共得到 {STATS['total_tokens']:,} 个有效词 token,"
        f"{STATS['unique_words']:,} 个不重复词类。",
    ]
    for s in aux:
        add_p(doc, s)


def build_ch3_methods(doc):
    """第三章 数据分析与可视化(方法步骤)。"""
    add_heading(doc, "三、数据分析与可视化", level=0)
    add_p(doc,
        "本研究的分析框架对标 ROST CM6 软件的核心功能,但全部使用 "
        "Python 重新实现,以获得对分词词典、共现阈值、社区检测算法、"
        "情感判别模型等关键参数的完全控制。下面分别说明三类分析的"
        "具体步骤。")

    add_heading(doc, "3.1 文本分词及词频分析", level=1)
    add_p(doc,
        "(1) 使用 jieba.posseg 对每条评论做带词性的分词;(2) 按词性"
        "和长度过滤,得到有效词列表;(3) 用 collections.Counter 在"
        "整体、国内、国际三个层级上做词频统计;(4) 用 matplotlib"
        "绘制 Top50 词频条形图,用 wordcloud 库绘制词云图——整体词云"
        "采用熊猫主题黑白灰+竹叶绿配色,国内组采用红色系,国际组采用"
        "蓝绿色系;(5) 设计「差异词」分析:对国内 Top200 ∪ 国际 Top200"
        " 中的每个词,计算 国内频次 /(国际频次 + 5) 作为国内特征比、"
        "国际频次 /(国内频次 + 5) 作为国际特征比,定位出只在一组高频"
        "的"
        "「特征词」,从而精准识别两组游客的关注差异。")

    add_heading(doc, "3.2 社会语义网络分析", level=1)
    add_p(doc,
        "(1) 取整体 Top30 高频词作为统一节点集(保证国内/国际/整体"
        "三张网络节点一致、可直接对比);(2) 在评论文本上做共现窗口"
        "扫描——以一条评论为窗口,两个词共同出现则共现计数 +1;"
        "(3) 设置共现阈值 ≥ 25 过滤弱关系,只保留强语义共现;"
        "(4) 用 networkx 计算度中心性、紧密中心性、介数中心性,识别"
        "网络核心节点;(5) 用 Louvain 算法做社区检测,自动归并语义"
        "相近的词成主题社区;(6) 同时输出静态 PNG 网络图(matplotlib)"
        "和交互式 HTML 网络图(pyvis,可在浏览器中拖拽节点查看)。")

    add_heading(doc, "3.3 情感分析", level=1)
    add_p(doc,
        "传统情感词典(知网 HowNet、台大 NTUSD)对评论语境理解粗糙,"
        "且无法应对国际版机翻文本的语义偏移。本研究改用 Google "
        "Gemini 2.5 大模型 API 做篇章级情感判别——对每条评论返回"
        "{sentiment: 正面/中性/负面, intensity: 强/中/弱, aspects: 8 类方面"
        "标签, is_translated: 是否疑似机翻}等字段,既能捕捉细粒度"
        "情感,又能识别翻译噪音。具体流程:(1) 设计结构化 prompt,"
        "强制模型输出 JSON;(2) 批量调用 API 并做断点续传,失败重试 "
        "3 次;(3) 输出未加权的分层样本情感分布,用于国内/国际对比;"
        "(4) 再按清洗后总体的 group×score 单元格数量为样本赋权,"
        "估计更接近自然总体口碑的情感比例。")

    add_heading(doc, "3.4 主题维度提炼", level=1)
    add_p(doc,
        "在上述三类分析的基础上,本研究采用「LLM 辅助的演绎式编码」"
        "方法,以旅游学经典的「目的地形象 + 满意度维度」理论框架为"
        "脚手架,让 Gemini 阅读整体 Top50 高频词、分组 Top30、特征词、"
        "典型评论摘要和情感分布表,提炼出 5 个一级主题维度,并产出"
        "每个维度的国内/国际情感得分对比雷达图。")


def build_ch4_results(doc):
    """第四章 数据结果分析(报告核心)。"""
    add_heading(doc, "四、数据结果分析", level=0)

    # ===========================================================
    # 4.1 词频分析结果
    # ===========================================================
    add_heading(doc, "4.1 词频分析结果", level=1)
    add_p(doc,
        f"阶段 1 为降低机翻文本对高频词的干扰,在词频分析中剔除疑似机翻样本,"
        f"保留 {STATS['word_freq_n']} 条评论参与分词统计,共得到 "
        f"{STATS['total_tokens']:,} 个有效词 token,{STATS['unique_words']:,} 个"
        "不重复词。国际组"
        f"约 {STATS['pct_translated_intl']}% 评论疑似为机器翻译,翻译过程会导致词汇被"
        "「中性化」收敛。表 4-1 列出了整体词频前 50 名。")

    # Top50 表(整体)
    top50 = pd.read_csv(OUT_DIR / "1_word_freq" / "all_top50.csv")
    add_table_from_df(doc, top50,
                      caption="表 4-1  整体游客评论高频词 Top50",
                      col_widths_cm=[1.5, 4.5, 3, 3], font_size=10)

    dim_path = OUT_DIR / "4_themes" / "dimensions_table.csv"
    if dim_path.exists():
        dim_names = pd.read_csv(dim_path, encoding="utf-8-sig")["维度名称"].astype(str).tolist()
    else:
        dim_names = ["核心吸引物感知", "游览流程与运营效率", "环境景观与场所属性",
                     "服务质量与附加价值", "综合体验与情感评价"]
    top_terms = "、".join(
        f"{r['word']}({int(r['count'])})" for _, r in top50.head(8).iterrows()
    )
    add_p(doc,
        f"从整体词频看,排名前列的「{top_terms}」勾勒出游客评论的核心叙事——"
        "「熊猫本体」「游园体验」与「运营痛点」三股力量同时存在。结合"
        f"旅游学经典的目的地形象框架,本研究将 Top50 高频词整合归纳为"
        f"{len(dim_names)} 个主题维度:{'、'.join(dim_names)}。该维度框架"
        "不仅在词频层面有可观测的证据支撑,也将在 4.4 节的雷达分析中"
        "获得情感得分上的进一步印证。")

    add_image(doc, IMG_WC_ALL, width_cm=14,
              caption=f"图 4-1  整体评论词云图(词频样本 {STATS['word_freq_n']} 条)")

    add_p(doc,
        "进一步将词云按国内/国际分组绘制,可以更直观地看到关注点的"
        "结构性差异:国内组(图 4-2)中的「观光车、孩子、排队、讲解、"
        "南门、电瓶车、挂件」等词显著放大,体现出本土游客关心园区"
        "运营效率、亲子体验和文创消费;国际组(图 4-3)中的「公园"
        "(Park)、参观、很大、值得、护照、班车、巴士」等词显著放大,"
        "其中「护照、班车、巴士、访问、预订」等词带有明显的英文翻译"
        "痕迹,印证了国际样本以 Trip.com 机翻为主的事实。")

    add_image(doc, IMG_WC_DOM, width_cm=14,
              caption="图 4-2  国内游客评论词云(红色系)")
    add_image(doc, IMG_WC_INTL, width_cm=14,
              caption="图 4-3  国际游客评论词云(蓝绿色系)")

    add_p(doc,
        "为量化两组关注点差异,本研究进一步绘制了 Top30 词频金字塔"
        "对比图(图 4-4)以及特征词对比图(图 4-5)。前者以并集 Top30"
        "为共享纵轴,左右镜像展示两组词的频次对比;后者展示"
        "「只在一组高频」的特征词。")

    add_image(doc, IMG_COMPARISON, width_cm=15,
              caption="图 4-4  国内游客 vs 国际游客 高频词金字塔对比")
    add_image(doc, IMG_DISTINCT, width_cm=15,
              caption="图 4-5  国内 vs 国际游客的特征词对比")

    # 差异词关键表
    dist_dom = pd.read_csv(OUT_DIR / "1_word_freq" / "distinctive_domestic.csv").head(10)
    dist_intl = pd.read_csv(OUT_DIR / "1_word_freq" / "distinctive_intl.csv").head(10)

    def _feature_summary(df_, n=6):
        parts = []
        for _, r in df_.head(n).iterrows():
            parts.append(f"{r['word']}(比值 {float(r['比值']):.3g})")
        return "、".join(parts)

    add_p(doc,
        "表 4-2 和表 4-3 分别列出两组特征词 Top10。国内特征词以"
        f"「{_feature_summary(dist_dom)}」为代表,反映出本地游客更关心"
        "园内交通、观赏效率、明星熊猫和具体消费环节;国际特征词以"
        f"「{_feature_summary(dist_intl)}」为代表,显示国际游客更常从"
        "整体环境、游览便利性和翻译语境表达目的地印象。")

    add_table_from_df(doc, dist_dom,
                      caption="表 4-2  国内游客特征词 Top10",
                      col_widths_cm=[2.5, 2.4, 2.4, 2, 2, 2], font_size=10)
    add_table_from_df(doc, dist_intl,
                      caption="表 4-3  国际游客特征词 Top10",
                      col_widths_cm=[2.5, 2.4, 2.4, 2, 2, 2], font_size=10)

    # ===========================================================
    # 4.2 社会语义网络结果
    # ===========================================================
    add_heading(doc, "4.2 社会语义网络分析结果", level=1)
    add_p(doc,
        "本研究在整体 Top30 高频词上构建共现网络,共现阈值设置为 25,"
        "并使用 Louvain 算法做社区检测。图 4-6 展示了国内/国际两组的"
        "网络对比图,图中节点大小代表度中心性,节点颜色代表 Louvain "
        "算法划分的社区归属,边的粗细代表共现强度。")

    add_image(doc, IMG_NET_COMPARE, width_cm=15.5,
              caption="图 4-6  国内 vs 国际游客 社会语义网络对比")

    add_p(doc,
        "从网络拓扑结构看,两组网络都呈现「熊猫」单核辐射型,但次级"
        "结构差异巨大。表 4-4 列出了两组度中心性 Top10,可直观看到"
        "次级核心词的分化:")

    cent = pd.read_csv(OUT_DIR / "2_network" / "centrality_comparison.csv")
    add_table_from_df(doc, cent,
                      caption="表 4-4  国内/国际游客网络 度中心性 Top10 对比",
                      col_widths_cm=[1.5, 3, 3.5, 3, 3.5], font_size=10.5)

    dom_core = "、".join(cent["国内词"].dropna().astype(str).head(10).tolist())
    intl_core = "、".join(cent["国际词"].dropna().astype(str).head(10).tolist())
    add_p(doc,
        f"国内组核心圈层为「{dom_core}」,带有强烈的具体场景与运营痛点色彩;"
        f"国际组核心圈层则为「{intl_core}」,呈现出更宏观、概括性的游览印象。"
        "Louvain 社区检测结果亦显示出结构差异:国内游客评论叙事更碎片化、"
        "关注更具体;国际游客评论叙事更概括、聚焦在整体印象。")

    # ===========================================================
    # 4.3 情感分析结果
    # ===========================================================
    add_heading(doc, "4.3 情感分析结果", level=1)
    add_p(doc,
        "本研究使用 Google Gemini 2.5 对全部 2000 条评论做篇章级情感"
        "判别。需再次强调:本研究的 2000 条样本为评分分层抽样产物,"
        "国内组 1 星配额 300 条(总体中 1 星仅占约 23%),国际组 5 星"
        "配额 470 条(总体中 5 星约占 42%)。这种设计使得各评分层都有"
        "足够样本进行组间对比和文本分析,但未加权的情感比例不能直接"
        "解释为携程总体自然口碑。因此本节同时报告两套口径的结果:")
    add_p(doc,
        "口径一:「分层样本(未加权)」——直接统计 2000 条样本中的情感"
        "分布,适用于国内/国际对比分析,回答「两组游客在相同评分层内"
        "的情感表达有何差异」。")
    add_p(doc,
        "口径二:「按总体评分分布加权」——以清洗后全量数据(7023 条,"
        "时间范围与样本一致)中每个 group×score 单元格的实际数量除以"
        "样本中同单元格数量作为权重,还原各层在总体中的真实比例,"
        "用于估计「如果对所有游客做情感分析,整体口碑大致如何」。")

    # 情感分段统计表(从 sentiment_results.csv 现场聚合,避免依赖未存在的 csv)
    sent_df = pd.read_csv(OUT_DIR / "3_sentiment" / "sentiment_results.csv")
    # 自动找列名
    sent_col = "sentiment" if "sentiment" in sent_df.columns else (
        "情感" if "情感" in sent_df.columns else sent_df.columns[1])
    group_col = "group"
    def _dist(sub):
        n = len(sub)
        if n == 0:
            return {"正面": "0 (0.0%)", "中性": "0 (0.0%)", "负面": "0 (0.0%)"}
        c = sub[sent_col].value_counts()
        return {k: f"{int(c.get(k, 0))} ({c.get(k, 0)/n*100:.1f}%)" for k in ["正面", "中性", "负面"]}
    rows = []
    n_all = len(sent_df)
    n_dom = int((sent_df[group_col] == "国内游客").sum())
    n_int = int((sent_df[group_col] == "国际游客").sum())
    for label, sub in [(f"整体分层样本({n_all})", sent_df),
                       (f"国内游客分层样本({n_dom})", sent_df[sent_df[group_col]=="国内游客"]),
                       (f"国际游客分层样本({n_int})", sent_df[sent_df[group_col]=="国际游客"])]:
        d = _dist(sub)
        rows.append({"分组": label, "正面": d["正面"], "中性": d["中性"], "负面": d["负面"]})
    sent_table = pd.DataFrame(rows)
    add_table_from_df(doc, sent_table,
                      caption="表 4-5  分层样本情感分段统计结果(未加权)",
                      col_widths_cm=[4, 3.5, 3.5, 3.5], font_size=11)

    weighted_path = OUT_DIR / "3_sentiment" / "sentiment_weighted_distribution.csv"
    weighted_table = None
    if weighted_path.exists():
        weighted = pd.read_csv(weighted_path, encoding="utf-8-sig")
        weighted_table = pd.DataFrame({
            "分组": weighted["分组"],
            "估计总体量": weighted["估计总体量"].map(lambda x: f"{float(x):.0f}"),
            "正面": weighted.apply(lambda r: f"{r['正面估计量']:.0f} ({r['正面占比']:.1f}%)", axis=1),
            "中性": weighted.apply(lambda r: f"{r['中性估计量']:.0f} ({r['中性占比']:.1f}%)", axis=1),
            "负面": weighted.apply(lambda r: f"{r['负面估计量']:.0f} ({r['负面占比']:.1f}%)", axis=1),
        })
        add_table_from_df(doc, weighted_table,
                          caption="表 4-6  按清洗后总体 group×score 分布加权的情感估计",
                          col_widths_cm=[4.5, 2.2, 3, 3, 3], font_size=10.5)

    add_p(doc,
        "表 4-5 说明在分析型分层样本中,国内游客组负面比例显著偏高,"
        "国际游客组正面比例显著占优,这一对比适合解释两组体验关注点"
        "和表达方式的差异。表 4-6 则把每个 group×score 单元格还原到"
        "清洗后总体中的实际规模,用于估计更接近自然总体口碑的情感比例。"
        "两套口径的含义不同:未加权结果突出对比,加权结果用于总体描述。"
        f"此外,国际组约 {STATS['pct_translated_intl']}% 评论存在机器翻译,"
        "翻译过程可能磨平原文中的情绪强度,因此国际组情感比例仍需结合"
        "翻译质量报告谨慎解释。")

    add_image(doc, IMG_SENT_WEIGHTED, width_cm=14,
              caption="图 4-7  按总体评分分布加权后的情感分布")
    add_image(doc, IMG_SENT_GROUP, width_cm=14,
              caption="图 4-8  国内 vs 国际游客 情感分布对比(分层样本)")
    add_image(doc, IMG_SENT_ASPECT, width_cm=15,
              caption="图 4-9  情感得分按维度的国内/国际对比")
    add_image(doc, IMG_SENT_TREND, width_cm=15,
              caption="图 4-10  评论情感随时间的月度趋势")

    # ===========================================================
    # 4.4 主题维度提炼
    # ===========================================================
    add_heading(doc, "4.4 主题维度提炼(综合分析)", level=1)
    add_p(doc,
        "综合 4.1~4.3 节的发现,本研究将熊猫基地的目的地形象提炼为"
        "五个一级主题维度(表 4-7),并计算每个维度的国内、国际"
        "情感得分,绘制雷达图(图 4-11)。")

    dim_df = pd.read_csv(OUT_DIR / "4_themes" / "dimensions_table.csv")
    # 把"包含的高频词"和"典型评论"列缩短一点,避免太宽
    if "包含的高频词" in dim_df.columns:
        dim_df["包含的高频词"] = dim_df["包含的高频词"].str.replace("、", "、").str[:60] + "…"
    if "典型评论" in dim_df.columns:
        dim_df["典型评论"] = dim_df["典型评论"].str.replace("\n", " ").str[:40] + "…"
    # 只挑核心列
    keep_cols = [c for c in ["维度名称","覆盖评论数","整体情感得分","国内情感得分","国际情感得分"]
                 if c in dim_df.columns]
    dim_df_compact = dim_df[keep_cols]
    add_table_from_df(doc, dim_df_compact,
                      caption="表 4-7  五大主题维度及其情感得分",
                      col_widths_cm=[5, 2.5, 2.8, 2.8, 2.8], font_size=10.5)

    add_image(doc, IMG_RADAR, width_cm=12,
              caption="图 4-11  五大维度国内/国际情感得分雷达图")

    max_gap_row = dim_df.loc[
        (dim_df["国内情感得分"].astype(float)
         - dim_df["国际情感得分"].astype(float)).abs().idxmax()
    ]
    max_gap = abs(float(max_gap_row["国内情感得分"])
                  - float(max_gap_row["国际情感得分"]))
    add_p(doc,
        "雷达图最显著的特征是国内组和国际组的情感得分几乎处于"
        "「镜像对立」状态:在五个维度上,国内得分整体偏负、国际得分"
        "整体偏正。当前数据中跨组差距最大的维度为"
        f"「{max_gap_row['维度名称']}」(国内 {float(max_gap_row['国内情感得分']):+.3f} / "
        f"国际 {float(max_gap_row['国际情感得分']):+.3f},差距 {max_gap:.3f})。"
        "这意味着,即使在面对同一个景点的同一群熊猫时,国内与国际游客的"
        "「期望—感知」框架截然不同:国内游客以本土同价位景区为对照"
        "锚点,容忍度低、批判敏感;国际游客以「中国旅行的稀缺体验」"
        "为对照锚点,容忍度高、整体倾向正面。")


def build_ch5_conclusion(doc):
    add_heading(doc, "五、结论与建议", level=0)

    add_heading(doc, "5.1 主要发现", level=1)
    findings = [
        f"(1) 核心吸引力突出,但运营挑战显著。「熊猫」以 "
        f"{STATS['word_counts'].get('熊猫', 0)} 次的绝对高频词位居首位,"
        "证实了其作为核心吸引物的地位。然而,"
        f"「排队」({STATS['word_counts'].get('排队', 0)} 次)、"
        f"「拥挤」({STATS['word_counts'].get('拥挤', 0)} 次)、"
        f"「看不到」({STATS['word_counts'].get('看不到', 0)} 次)等"
        "负面感知词广泛存在,表明基地在游客流量管理和内部运营方面"
        "存在普遍挑战。",
        "(2) 国内外游客感知呈现极端两极分化。国内游客对基地的整体"
        "评价偏向负面,尤其在「性价比」「餐饮购物」「票务交通」方面"
        "表现出强烈不满;相比之下,国际游客整体满意度较高,对上述方面"
        "的评价均持正面态度。",
        "(3) 关注焦点与期望存在根本差异。国内游客高度关注运营效率"
        "(如「观光车」「电瓶车」)、特定明星熊猫(如「花花」)及"
        "具体的场馆体验,并对「不合理」现象敏感;国际游客则更侧重"
        "基地作为「公园」的整体环境、自然美感(如「美丽」)以及入园和"
        "交通的便利性(如「护照」「预订」)。",
    ]
    for f in findings:
        add_p(doc, f)

    add_heading(doc, "5.2 理论与方法贡献", level=1)
    add_p(doc,
        "本研究通过应用基于方面的情感分析(ABSA)方法对海量用户生成"
        "内容(UGC)进行深度挖掘,为目的地形象理论提供了实证支持。它"
        "不仅确认了目的地形象的多维度特性,更突出地揭示了文化背景"
        "(国内外游客)作为关键调节变量,如何显著影响游客对目的地"
        "具体形象维度(如核心吸引物、服务设施、环境景观)的情感感知"
        "和整体满意度。尤其在「性价比」这一综合价值感知维度上,国内外"
        f"游客在「{STATS['max_dim_gap_name']}」维度上高达 "
        f"{STATS['max_dim_gap']} 的情感差异,为跨文化旅游研究中游客期望与"
        "感知价值的构建提供了新的观察视角。")

    add_heading(doc, "5.3 管理启示与建议", level=1)
    advice = [
        "(1) 针对国内游客优化运营效率与提升价值感知。鉴于国内游客对"
        "「观光车」「排队」和「性价比」的强烈不满,基地应重点投资于"
        "智能分流系统、增加观光车班次、优化排队管理流程,并提供更具"
        "吸引力且物有所值的餐饮、购物及增值服务(如特色文创、互动体验),"
        "以直接回应其对效率和价值的期望。",
        "(2) 强化国际游客的无障碍体验。国际游客整体满意度高,但对"
        "「护照」验证和「预订」流程的关注提示了潜在痛点。基地应进一步"
        "简化国际游客入园流程,提供多语言的清晰指示牌、导览信息和工作"
        "人员服务,确保国际游客能够顺畅便捷地享受其所珍视的「公园」体验。",
        "(3) 实施差异化的市场沟通策略。针对国内市场,应强调基地的"
        "管理改进、特定明星熊猫的观赏攻略和高性价比的游览方案。而对"
        "国际市场,则应突出大熊猫的独特性、基地作为自然保护区的生态"
        "价值以及整体轻松愉悦的「公园」式体验。",
    ]
    for a in advice:
        add_p(doc, a)

    add_heading(doc, "5.4 研究局限", level=1)
    add_p(doc,
        "本研究主要基于在线评论数据,可能存在样本选择偏差:即发布评论"
        "的游客往往是体验特别好或特别差的群体,未能全面反映所有游客"
        "的真实感受。此外,「国际游客」作为一个集合概念,涵盖了多种"
        f"文化背景和语言群体,且约 {STATS['pct_translated_intl']}% 的国际样本为机器翻译文本"
        f"(国际组占 {STATS['pct_translated_intl']}%),机翻过程可能让情感判断偏中性化。未来研究"
        "可进一步细分国际游客群体,并采用机翻识别+人工抽检的方法,以"
        "揭示更深层次的文化差异对目的地形象感知的影响。")


# ============================================================
# 四、主流程
# ============================================================
def main():
    # 切换到 panda_analysis(此脚本所在目录)以保证相对路径正确
    here = Path(__file__).parent
    os.chdir(here)

    # 加载动态统计数据(必须在 chdir 之后,因为路径是相对的)
    global STATS
    STATS = _load_stats()
    print(f"[stats] tokens={STATS['total_tokens']:,}  "
          f"unique={STATS['unique_words']:,}  "
          f"translated={STATS['pct_translated']}%")

    doc = Document()

    # 全局页面设置:A4,2.5cm 边距
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 页脚页码
    add_page_number_footer(doc)

    # 默认 Normal 样式
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    # 1. 封面
    build_cover(doc)
    add_page_break(doc)

    # 2. 第一章
    build_ch1_collection(doc)
    add_page_break(doc)

    # 3. 第二章
    build_ch2_preprocess(doc)
    add_page_break(doc)

    # 4. 第三章
    build_ch3_methods(doc)
    add_page_break(doc)

    # 5. 第四章
    build_ch4_results(doc)
    add_page_break(doc)

    # 6. 第五章
    build_ch5_conclusion(doc)

    # 保存
    doc.save(str(REPORT_PATH))
    size_kb = os.path.getsize(REPORT_PATH) / 1024
    print(f"\n✅ 报告已生成: {REPORT_PATH}")
    print(f"   大小: {size_kb:.1f} KB")
    print(f"   完整路径: {Path.cwd() / REPORT_PATH}")


if __name__ == "__main__":
    main()
